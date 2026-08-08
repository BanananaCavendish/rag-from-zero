"""生成虚构企业知识语料,覆盖 4 种格式:PDF / Word(docx) / Markdown / HTML。

虚构公司:晨光科技有限公司。语料 5 份:
  - employee_handbook.md        员工手册
  - travel_reimbursement.pdf    差旅报销制度
  - product_user_guide.docx     晨光云文档用户手册
  - faq.html                    常见问题 FAQ
  - security_policy.md          信息安全管理制度

用法:python scripts/generate_corpus.py
输出:data/corpus/ 下的 5 个文件。

说明:
  - 内容刻意包含具体数字 / 政策编号 / 专有名词,用于验证「精确词检索」(BM25 强)
    与「语义检索」(向量强)的分工。
  - 所有内容完全虚构,仅供学习演示。
"""

import sys
from pathlib import Path

# 让脚本无论从哪个目录运行都能 import backend
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
# Windows 控制台默认 GBK,打印 emoji/生僻字会崩;统一重配置为 UTF-8
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

from backend.core import config

CORPUS_DIR = config.CORPUS_DIR

# =====================================================================
# 内容定义(每一节:标题, 正文)
# =====================================================================

EMPLOYEE_HANDBOOK = [
    ("第一章 总则",
     "本手册适用于晨光科技有限公司全体正式员工及实习生。所有员工应熟悉并遵守本手册的各项规定,"
     "如与劳动合同另有约定,以劳动合同为准。"),
    ("第二章 工作时间",
     "公司实行每周五天工作制,工作时间为上午 9:00 至下午 18:00,午休 12:00 至 13:30。"
     "为方便通勤,允许员工弹性上下班,弹性范围为前后半小时。每日需在考勤系统打卡。"),
    ("第三章 试用期",
     "新员工试用期为 3 个月,表现优秀者可提前转正,但试用期不得少于 1 个月。"
     "试用期内如不符合录用条件,公司可解除劳动合同。"),
    ("第四章 年休假",
     "入职满 1 年享受年休假 5 天;满 3 年享受 8 天;满 5 年享受 12 天,上限 15 天。"
     "年休假原则上在当年度内休完,逾期未休且未安排调休的,视为自动放弃。"),
    ("第五章 病假与医疗",
     "员工请病假需提供三级甲等医院的诊断证明或病假条。病假期间工资按当地规定执行,"
     "医疗期累计不超过国家规定。"),
    ("第六章 加班与调休",
     "公司提倡高效工作,不鼓励长期加班。确因工作需要加班的,需提前填写加班审批单。"
     "工作日加班可申请等时调休,周末加班按劳动法支付加班费。"),
    ("第七章 离职办理",
     "员工离职需提前 30 天书面通知公司。离职前需完成工作交接、归还办公设备,"
     "并结清各类借款与报销。最后一个工作日凭《离职交接单》办理离职手续。"),
]

TRAVEL_REIMBURSEMENT = [
    ("一、适用范围",
     "本制度适用于公司全体员工因公出差产生的交通、住宿、餐饮等费用的报销。"
     "出差须提前在 OA 系统提交出差申请,经部门负责人审批后方可出行。"),
    ("二、报销所需材料",
     "报销时必须提供以下材料:(一)合规的增值税发票;(二)行程单或登机牌、车票等凭证;"
     "(三)经审批的《出差审批单》;(四)如有超标项,需另附特别审批说明。材料不全不予受理。"),
    ("三、交通费用标准",
     "高铁/动车:原则上乘坐二等座,单程超过 4 小时可申请一等座。"
     "飞机:经济舱,机票需提前 3 个工作日通过公司差旅平台预订。"
     "市内交通:地铁、公交、打车合计每日上限 100 元。"),
    ("四、住宿费用标准",
     "一线城市(北上广深)住宿标准为每晚 500 元,其他城市每晚 350 元。"
     "超出标准的部分需在出差前获得部门负责人特别审批,否则超出部分自理。"),
    ("五、餐饮补贴",
     "出差期间按日发放餐饮补贴,标准为每日 100 元,不再另行报销餐费发票。"
     "出差不足半日的按半日计发。"),
    ("六、报销流程与时限",
     "出差结束后应在 3 个工作日内提交报销申请。流程为:OA 提交 → 部门负责人审批 → "
     "财务审核 → 出纳打款。审核通过后 15 个工作日内到账。")
]

PRODUCT_USER_GUIDE = [
    ("产品简介",
     "晨光云文档是公司自研的企业级在线文档协作平台,支持在线编辑、多人实时协作、"
     "版本历史与精细化权限管理,是团队知识沉淀的统一入口。"),
    ("在线编辑",
     "新建文档支持富文本、Markdown、表格、图片与代码块。文档自动实时保存,"
     "编辑器支持离线模式,断网时自动保存到本地缓存,恢复联网后自动同步。"),
    ("多人协作",
     "可邀请成员以「可查看」「可评论」「可编辑」三种权限协作同一文档。"
     "多人同时编辑时,系统会以颜色区分不同成员的编辑内容,并支持评论与 @ 提醒。"),
    ("版本历史",
     "文档每 10 分钟自动生成一个版本快照,最多保留 100 个版本。"
     "可从版本历史中查看任意历史版本,支持一键回滚,误删内容可随时找回。"),
    ("导出与分享",
     "文档支持导出为 PDF、Word 与 Markdown 格式。分享时可为链接设置有效期、"
     "访问密码与访问人数上限,外部链接默认只读。"),
    ("常见故障处理",
     "如果文档无法保存,请先检查网络连接,确认右下角同步状态图标为绿色;"
     "若提示权限不足,请联系文档所有者或系统管理员调整权限;"
     "若页面白屏,请强制刷新或清除浏览器缓存后重试。")
]

FAQ = [
    ("如何重置账号密码?",
     "在公司门户登录页点击「忘记密码」,输入绑定的企业邮箱或手机号,"
     "系统将发送重置链接。链接 30 分钟内有效,重置后需重新登录。"),
    ("如何开通企业邮箱?",
     "向部门主管提出申请,由行政部统一开通。开通后使用工号作为账号前缀,"
     "初始密码为身份证后六位,首次登录必须修改。"),
    ("如何申请 VPN 远程办公?",
     "在 IT 服务台提交 VPN 申请,注明用途与预计使用期限。"
     "审批通过后 IT 将发放一次性激活码,一个激活码仅限一人使用,禁止共享。"),
    ("如何预订会议室?",
     "通过 OA 系统的会议室模块预订,支持按容量与设备筛选。"
     "预订后如需取消,请至少提前 30 分钟释放资源,方便他人使用。"),
    ("出差报销最长可以多久?",
     "出差结束后应在 3 个工作日内提交报销申请,审核通过后 15 个工作日内到账。"
     "逾期超过 30 天提交的报销申请需部门负责人书面说明原因。"),
    ("如何申请加班审批?",
     "填写加班审批单,注明加班日期、时长与原因,由部门负责人审批。"
     "工作日加班可申请等时调休,周末加班按劳动法支付加班费。"),
]

SECURITY_POLICY = [
    ("一、口令管理",
     "员工账号口令必须不少于 12 位,并同时包含大小写字母、数字与特殊字符。"
     "口令每 90 天必须更换一次,不得与历史 5 次口令重复,禁止在多个系统间复用同一口令。"),
    ("二、数据分级",
     "公司数据按敏感程度分为三级:公开(对外宣传资料)、内部(员工通讯录、内部制度)、"
     "机密(客户数据、财务数据、源代码、商业计划)。机密数据禁止外发、禁止存储于个人设备。"),
    ("三、机密数据处理",
     "处理机密数据必须使用公司发放的加密笔记本或经 IT 批准的沙箱环境。"
     "打印机密文件需登记,废弃的机密文件必须用碎纸机销毁。"),
    ("四、终端与网络",
     "公司电脑必须安装统一的安全软件并保持自动更新。禁止在未经批准的设备上处理机密数据,"
     "禁止使用个人 VPN 绕过公司网络出口。"),
    ("五、安全事件上报",
     "发现疑似钓鱼邮件、异常登录、数据泄露等安全事件,必须在 1 小时内上报 IT 安全团队,"
     "并保留现场证据,不得自行删除相关日志。"),
]

# =====================================================================
# 各格式生成
# =====================================================================


def _md(filename: str, title: str, sections: list[tuple[str, str]]) -> Path:
    lines = [f"# {title}", ""]
    for heading, text in sections:
        lines += [f"## {heading}", "", text, ""]
    path = CORPUS_DIR / filename
    path.write_text("\n".join(lines), encoding="utf-8")
    print(f"✅ {filename} ({len(sections)} 节)")
    return path


def _pdf(filename: str, title: str, sections: list[tuple[str, str]]) -> Path:
    """用 reportlab 内置 STSong-Light CID 字体渲染中文 PDF,跨平台无需外部字体。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    base = getSampleStyleSheet()
    title_style = ParagraphStyle("CnTitle", parent=base["Title"], fontName="STSong-Light",
                                 fontSize=18, leading=24)
    heading_style = ParagraphStyle("CnHeading", parent=base["Heading2"], fontName="STSong-Light",
                                   fontSize=13, leading=20)
    body_style = ParagraphStyle("CnBody", parent=base["BodyText"], fontName="STSong-Light",
                                fontSize=10.5, leading=18)

    story = [Paragraph(title, title_style), Spacer(1, 16)]
    for heading, text in sections:
        story.append(Paragraph(heading, heading_style))
        story.append(Paragraph(text, body_style))
        story.append(Spacer(1, 10))

    path = CORPUS_DIR / filename
    SimpleDocTemplate(str(path), pagesize=A4, title=title).build(story)
    print(f"✅ {filename} ({len(sections)} 节)")
    return path


def _docx(filename: str, title: str, sections: list[tuple[str, str]]) -> Path:
    import docx

    doc = docx.Document()
    doc.add_heading(title, 0)
    for heading, text in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(text)
    path = CORPUS_DIR / filename
    doc.save(path)
    print(f"✅ {filename} ({len(sections)} 节)")
    return path


def _html(filename: str, title: str, sections: list[tuple[str, str]]) -> Path:
    """手写语义化 HTML,BSHTMLLoader 可直接抽取正文。"""
    body_items = "".join(
        f"<h2>{h}</h2><p>{t}</p>" for h, t in sections
    )
    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head><meta charset="utf-8"><title>{title}</title></head>
<body>
<h1>{title}</h1>
{body_items}
</body>
</html>"""
    path = CORPUS_DIR / filename
    path.write_text(html, encoding="utf-8")
    print(f"✅ {filename} ({len(sections)} 节)")
    return path


def main() -> None:
    CORPUS_DIR.mkdir(parents=True, exist_ok=True)
    _md("employee_handbook.md", "晨光科技员工手册", EMPLOYEE_HANDBOOK)
    _pdf("travel_reimbursement.pdf", "晨光科技差旅报销制度(试行)", TRAVEL_REIMBURSEMENT)
    _docx("product_user_guide.docx", "晨光云文档用户手册", PRODUCT_USER_GUIDE)
    _html("faq.html", "晨光科技常见问题 FAQ", FAQ)
    _md("security_policy.md", "晨光科技信息安全管理制度", SECURITY_POLICY)
    print(f"\n🎉 已生成 {len(list(CORPUS_DIR.glob('*')))} 份语料 → {CORPUS_DIR}")


if __name__ == "__main__":
    main()
